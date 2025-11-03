# sports_team/report.py
import os
from docx import Document
from sports_team.team import Team
from sports_team.db import get_team_match_stats
from sports_team.utils import timed

@timed
def save_team_report_docx(team: Team, filename: str):
    """Создаёт отчёт о команде и сохраняет его в формате .docx"""
    report_dir = os.path.join(os.path.dirname(__file__), "..", "report")
    os.makedirs(report_dir, exist_ok=True)
    filepath = os.path.abspath(os.path.join(report_dir, filename))

    doc = Document()
    doc.add_heading(f"Отчёт о команде: {team.name}", level=1)

    # --- Общая статистика ---
    total_goals = team.total_goals()
    total_games = getattr(team, "total_games", lambda: 0)()
    player_count = len(team.players)
    
    doc.add_paragraph(f"Количество игроков: {player_count}")
    doc.add_paragraph(f"Общие голы: {total_goals}")
    
  
    # --- Статистика матчей ---
    doc.add_heading("Статистика матчей", level=2)
    stats = get_team_match_stats(team.name)
    doc.add_paragraph(f"Сыграно матчей: {stats['Матчи']}")
    doc.add_paragraph(f"Побед: {stats['Победы']}")
    doc.add_paragraph(f"Поражений: {stats['Поражения']}")
    doc.add_paragraph(f"Ничьих: {stats['Ничьи']}")

    # --- Таблица игроков ---
    doc.add_heading("Состав команды", level=2)
    table = doc.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Имя"
    hdr_cells[1].text = "Номер"
    hdr_cells[2].text = "Позиция"
    hdr_cells[3].text = "Матчи"
    hdr_cells[4].text = "Голы"
    hdr_cells[5].text = "Передачи"

    for player in team.players:
        row = table.add_row().cells
        row[0].text = str(player.name)
        row[1].text = str(player.number)
        row[2].text = str(getattr(player, "position", None) or player.role())
        row[3].text = str(player.games)
        row[4].text = str(player.goals)
        row[5].text = str(player.assists)

    # --- Лучший бомбардир ---
    if team.players:
        top = max(team.players, key=lambda p: p.goals)
        if top.goals > 0:
            doc.add_paragraph(f"Лучший бомбардир: {top.name} ({top.goals} голов)")

    # --- Сохранение и открытие ---
    doc.save(filepath)
    print(f"📄 Отчёт сохранён: {filepath}")
    try:
        os.startfile(filepath)
    except Exception:
        pass
